#!/usr/bin/env python3
"""Generează un .docx cu tabelele de interes pentru Raportul 3 (Componenta Flow)."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# font implicit
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def add_table(title, headers, rows, note=None):
    doc.add_heading(title, level=2)
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    # header
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
    # body
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ""
            run = cells[j].paragraphs[0].add_run(str(val))
            # îngroașă rândurile-cheie marcate cu *
            if str(val).startswith("*") :
                run.text = str(val).strip("*")
                run.bold = True
    if note:
        p = doc.add_paragraph()
        r = p.add_run(note)
        r.italic = True
        r.font.size = Pt(9)
    doc.add_paragraph()


# ---- titlu ----
h = doc.add_heading("Componenta Flow — tabele de rezultate (Raportul 3)", level=0)

p = doc.add_paragraph()
r = p.add_run("Detector hibrid XGBoost (supervised) + Autoencoder (unsupervised), "
              "evaluat pe seturile BCCC și ITU. Cifre verificate din artefactele de "
              "antrenare și evaluare.")
r.italic = True
doc.add_paragraph()

# ---- Tabel 1 ----
add_table(
    "Tabelul 1 — Seturile de date",
    ["Caracteristică", "BCCC Cloud DDoS 2024", "ITU 2023 (Kubernetes)"],
    [
        ["Flow records", "~700k", "~3,2M"],
        ["Features (ML)", "317", "79"],
        ["Tip atac", "DDoS (TCP)", "CVE-uri Kubernetes"],
        ["Dezechilibru (benign:atac)", "~1,44 : 1", "~10,6 : 1"],
        ["Holdout structural", "held-out day (temporal)", "leave-heavy-hitter-out (sampling)"],
        ["Rol", "set principal", "set comparativ"],
    ],
)

# ---- Tabel 2 ----
add_table(
    "Tabelul 2 — Hiperparametri",
    ["XGBoost (supervised)", "Autoencoder (unsupervised)"],
    [
        ["n_estimators = 500", "arhitectură: input→128→32→8→32→128→input"],
        ["max_depth = 8", "bottleneck = 8"],
        ["learning_rate = 0,1", "optimizer = Adam, lr = 1e-3"],
        ["scale_pos_weight = dinamic (≈1,44 / ≈10,6)", "loss = MSE · batch = 4096"],
        ["early_stopping = 30", "epoci = 30 (max) · patience = 5"],
        ["objective = binary:logistic · tree_method = hist", "prag = percentila 95 a MSE pe val benign"],
        ["seed = 42", "antrenat doar pe benign · StandardScaler · seed 42"],
    ],
)

# ---- Tabel 3 ----
add_table(
    "Tabelul 3 (A) — Prăbușirea: generalizarea XGBoost (in-distribution vs. holdout structural)",
    ["Set", "Protocol", "ROC-AUC", "Recall", "F1"],
    [
        ["BCCC", "random (in-distribution)", "0,999", "0,979", "0,979"],
        ["BCCC", "held-out day (temporal)", "0,953", "0,775", "0,847"],
        ["ITU", "random (in-distribution)", "0,999", "0,991", "0,879"],
        ["*ITU", "*LHO (sampling)", "*0,639", "*0,048", "*0,089"],
    ],
    note="Recall și F1 la pragul implicit 0,5. Colapsul ITU LHO (0,991 → 0,048) = signature coupling.",
)

# ---- Tabelul 4 (B) — fuziune pe holdout-uri oneste ----
add_table(
    "Tabelul 4 (B) — Rolul autoencoderului: câștigul fuziunii pe holdout-urile oneste",
    ["Strategie", "BCCC: recall@FPR=1%", "BCCC: ROC-AUC",
     "ITU LHO: recall@FPR=1%", "ITU LHO: ROC-AUC"],
    [
        ["XGBoost singur", "0,782", "0,972", "0,053", "0,659"],
        ["Autoencoder singur", "0,006", "0,718", "0,063", "0,857"],
        ["*Hybrid (weighted_70_30)", "*0,799", "*0,973", "*0,094", "*0,857"],
    ],
    note="Doar holdout-uri structurale (relevant pentru deployment; val omis intenționat). "
         "Câștig recall@FPR=1%: BCCC +1,7 pp, ITU LHO +4,1 pp. Pe ITU LHO fuziunea recuperează "
         "discriminarea (ROC-AUC 0,659 → 0,857) exact unde XGBoost se prăbușește prin signature coupling. "
         "Cifrele LHO sunt pe rânduri aliniate XGB/AE (comparație apples-to-apples).",
)

out = os.path.join(os.path.dirname(__file__), "tabele_raport3.docx")
doc.save(out)
print("Saved:", out)
