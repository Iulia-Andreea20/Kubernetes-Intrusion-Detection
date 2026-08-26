#!/usr/bin/env python3
"""Tabelul 1 + Tabelul 2 — performanța clasificatorului pe atacatori sintetici
și efectul testării tool-disjuncte cu unelte externe.
Format adaptiv pentru copy-paste (autofit + lățimi procentuale)."""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/iulia-andreeagrigore/Projects/Kubernetes-Intrusion-Detection/docs/thesis-ro/tabel1_actori_sintetici.docx"

doc = Document()

# Portrait A4 + margini moderate
section = doc.sections[0]
section.left_margin = Cm(2.0)
section.right_margin = Cm(2.0)
section.top_margin = Cm(2.0)
section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=40, bottom=40, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_width_pct(cell, pct_x50):
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(pct_x50))
    tcW.set(qn('w:type'), 'pct')
    tcPr.append(tcW)


def set_cell_shading(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def make_adaptive_table(headers, rows, col_pct, center_cols=None, mono_cols=None,
                        highlight_rows=None):
    if center_cols is None: center_cols = []
    if mono_cols is None: mono_cols = []
    if highlight_rows is None: highlight_rows = {}
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.autofit = True
    tblPr = t._tbl.tblPr
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    col_50ths = [int(round(p * 50)) for p in col_pct]

    # header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ''
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        if i in center_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        set_cell_border(c)
        set_cell_padding(c)
        set_cell_width_pct(c, col_50ths[i])
        set_cell_shading(c, 'D9D9D9')

    # data rows
    for ri, row in enumerate(rows, 1):
        for i, val in enumerate(row):
            c = t.rows[ri].cells[i]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            if i in center_cols:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            if i in mono_cols:
                r.font.name = 'Consolas'
                r.font.size = Pt(9.5)
            # highlight (de ex. cele cu drop mare)
            if (ri-1) in highlight_rows:
                set_cell_shading(c, highlight_rows[ri-1])
            set_cell_border(c)
            set_cell_padding(c)
            set_cell_width_pct(c, col_50ths[i])
    return t


# ============================================================
# TABEL 1 — 4 atacatori sintetici originali
# ============================================================
h = doc.add_heading(
    'Tabelul 1 — Performanța clasificatorului pe cei 4 atacatori sintetici originali (in-distribution)',
    level=2
)
h.runs[0].font.size = Pt(11)
h.paragraph_format.space_before = Pt(0)
h.paragraph_format.space_after = Pt(6)

HEADERS_1 = ['Actor sintetic', 'Tactică MITRE ATT&CK', 'N ferestre atac', 'Recall', 'FPR', 'AUC']
COL_PCT_1 = [18, 38, 14, 10, 10, 10]
CENTER_1  = [2, 3, 4, 5]

ROWS_1 = [
    ('victim-sa',          'Credential Access (T1078 / T1552.007)', '121', '100,0%', '0,66%', '1,000'),
    ('adversary-external', 'Privilege Escalation (T1078)',          '152', '100,0%', '0,66%', '1,000'),
    ('adversary-insider',  'Privilege Escalation — low-and-slow',   '246', '98,0%',  '0,66%', '0,999'),
    ('recon-sa',           'Discovery (T1069)',                     '234', '0,9%',   '0,66%', '0,264'),
]

make_adaptive_table(HEADERS_1, ROWS_1, COL_PCT_1, center_cols=CENTER_1, mono_cols=[0])

# Nota explicativă pentru Tabel 1
nota1 = doc.add_paragraph()
nota1.paragraph_format.space_before = Pt(8)
nota1.paragraph_format.space_after = Pt(0)
r = nota1.add_run('Notă (Tabel 1): ')
r.bold = True; r.font.size = Pt(9)
r = nota1.add_run(
    'recall = TP / (TP + FN) pe ferestrele actorului; FPR = FP / (FP + TN) global pe partiția '
    'benignă de test (identic pentru toți actorii); AUC = ROC AUC al clasificatorului XGBoost. '
    'recon-sa apare cu performanță sub-aleatoare prin design: reconul a fost EXCLUS din '
    'antrenarea clasificatorului (delegat regulii recon cu allowlist) datorită ambiguității '
    'comportamentale intrinsece cu automatizările benigne.'
)
r.italic = True; r.font.size = Pt(9)

# ============================================================
# TABEL 2 — Efectul testării tool-disjuncte cu unelte externe
# ============================================================
sep = doc.add_paragraph()
sep.paragraph_format.space_before = Pt(18)
sep.paragraph_format.space_after = Pt(0)

h2 = doc.add_heading(
    'Tabelul 2 — Performanța pe unelte red-team externe (testare tool-disjunctă) cu diferența față de scenariul sintetic',
    level=2
)
h2.runs[0].font.size = Pt(11)
h2.paragraph_format.space_before = Pt(0)
h2.paragraph_format.space_after = Pt(6)

HEADERS_2 = ['Actor original', 'Tool extern echivalent', 'N atac', 'Recall extern', 'AUC extern', 'Δ Recall vs sintetic']
COL_PCT_2 = [16, 36, 9, 12, 11, 16]
CENTER_2  = [2, 3, 4, 5]

ROWS_2 = [
    ('victim-sa',          'Stratus k8s.credential-access.dump-secrets *', '85',   '95,3%',  '0,997', '−4,7 pp'),
    ('adversary-external', 'Peirates (InGuardians, MITRE S0683)',          '23',   '100,0%', '1,000', '0,0 pp ✓'),
    ('adversary-insider',  'lowslow (atac diluat, identity-disjunct) †',   '826',  '49,9%',  '0,958', '−48,1 pp'),
    ('recon-sa',           'rakkess (kubectl access-matrix, krew plugin)', '9.697', '0,0%',  '0,635', '−0,9 pp ‡'),
]

# Highlight: linia adversary-insider (drop catastrofal -48pp) și linia adversary-external (ZERO drop, dovada generalizare)
highlight = {
    1: 'D5E8D4',  # adversary-external — verde deschis (dovadă reușită)
    2: 'FCE4D6',  # adversary-insider — portocaliu deschis (drop major)
}

make_adaptive_table(HEADERS_2, ROWS_2, COL_PCT_2, center_cols=CENTER_2, mono_cols=[0],
                    highlight_rows=highlight)

# Note pentru Tabel 2
nota2 = doc.add_paragraph()
nota2.paragraph_format.space_before = Pt(8)
nota2.paragraph_format.space_after = Pt(2)
r = nota2.add_run('Note (Tabel 2): ')
r.bold = True; r.font.size = Pt(9)
r = nota2.add_run(
    'Coloana „Δ Recall vs sintetic" arată diferența procentuală absolută (pp = puncte procentuale) '
    'între recall-ul pe testul tool-disjunct extern și recall-ul pe scenariul sintetic corespondent '
    '(Tabel 1). Reducerea catastrofală de 48 pp pentru adversary-insider expune limita reală a '
    'modelului față de atacurile diluate temporal (low-and-slow); generalizarea perfectă pe Peirates '
    '(0 pp drop) constituie dovada empirică anti-circularitate (Arp et al., 2022) — modelul prinde '
    'la fel de bine atacuri venind de la o unealtă red-team complet independentă.'
)
r.italic = True; r.font.size = Pt(9)

# Note cu marcajele *, †, ‡
nota3 = doc.add_paragraph()
nota3.paragraph_format.space_before = Pt(4)
nota3.paragraph_format.space_after = Pt(2)
nota3.paragraph_format.left_indent = Cm(0.4)
parts = [
    ('* ', True, 'Stratus dump-secrets este distribuit ca modul în suite-ul Stratus Red Team general (DataDog); nu există o unealtă terță dedicată exclusiv tacticii Credential Access pe planul de audit Kubernetes.'),
    ('\n† ', True, 'lowslow nu este o unealtă terță în sens strict, ci un sub-profil sintetic propriu cu identitate distinctă (adversary-stealth) și pattern temporal dilatat; conform cercetării sistematice (UNELTE_EXTERNE_VALIDARE.md), nu există unealtă terță pentru atacuri low-and-slow pe audit K8s.'),
    ('\n‡ ', True, 'Diferența de recall pe recon-sa este nesemnificativă semantic: ambele valori sunt aproape zero deoarece reconul nu este învățat de clasificator (delegat regulii recon prin design); pipeline-ul hibrid recuperează 100% pe rakkess prin regula recon cu allowlist.'),
]
for marker, italic, text in parts:
    r = nota3.add_run(marker)
    r.bold = True; r.font.size = Pt(8.5)
    r = nota3.add_run(text)
    r.italic = italic; r.font.size = Pt(8.5)

doc.save(OUT)
print(f"Saved: {OUT}")
