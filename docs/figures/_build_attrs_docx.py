#!/usr/bin/env python3
"""Genereaza tabele compacte cu atributele setului de date.

Strategia pt copy-paste corect:
- tblLayout=autofit (tabelul reflow-eaza la lipire in alt document)
- lățimi proportionale (procente, nu cm) -> se adaptează la latimea destinatiei
- tblW de 100% pe tabel
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/iulia-andreeagrigore/Projects/Kubernetes-Intrusion-Detection/docs/thesis-ro/atribute_set_date.docx"

doc = Document()

# A4 portrait + margini moderate (compatibil cu lucrarea ta tipica)
section = doc.sections[0]
section.left_margin = Cm(1.5)
section.right_margin = Cm(1.5)
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(9)


def set_cell_border(cell):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '808080')
        tcBorders.append(b)
    tcPr.append(tcBorders)


def set_cell_padding(cell, top=30, bottom=30, left=80, right=80):
    """Padding in twips (1/20 pt). Compact dar lizibil."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def set_cell_width_pct(cell, pct_x50):
    """Lățime celulă în 50-imi de procent (5000 = 100%). Asta permite reflow."""
    tcPr = cell._tc.get_or_add_tcPr()
    # sterg orice tcW existent
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(pct_x50))
    tcW.set(qn('w:type'), 'pct')
    tcPr.append(tcW)


def set_cell_shading(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_hex)
    cell._tc.get_or_add_tcPr().append(shd)


def make_adaptive_table(headers, rows, col_pct):
    """col_pct = lista de procente (insumeaza 100). Returneaza tabel adaptiv."""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    # Layout AUTOFIT — esential pt copy-paste
    t.autofit = True
    tblPr = t._tbl.tblPr
    # Sterg tblLayout existent si pun autofit
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'autofit')
    tblPr.append(tblLayout)
    # Latimea totala tabel = 100% din container
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '5000')   # 5000 = 100% (in 50-imi de procent)
    tblW.set(qn('w:type'), 'pct')
    tblPr.append(tblW)

    # Convertesc procentele in 50-imi (5000 = 100%)
    col_50ths = [int(round(p * 50)) for p in col_pct]

    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        c.text = ''
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        set_cell_border(c)
        set_cell_padding(c)
        set_cell_width_pct(c, col_50ths[i])
        set_cell_shading(c, 'D9D9D9')

    # Rânduri
    for ri, row in enumerate(rows, 1):
        for i, val in enumerate(row):
            c = t.rows[ri].cells[i]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(8.5)
            if i == 0:
                r.bold = True
            if i == 1:
                r.font.name = 'Consolas'
                r.font.size = Pt(8)
            set_cell_border(c)
            set_cell_padding(c)
            set_cell_width_pct(c, col_50ths[i])
    return t


# ============================================================
# TABEL 1 — 30 atribute folosite
# ============================================================
h1 = doc.add_heading('Tabelul 1 — Cele 30 atribute folosite de clasificator', level=1)
h1.runs[0].font.size = Pt(12)
h1.paragraph_format.space_before = Pt(0)
h1.paragraph_format.space_after = Pt(4)

# (#, atribut, calcul, semnificație) — procente care insumeaza 100
TBL1_HEADERS = ['#', 'Atribut', 'Calcul', 'Semnificație']
TBL1_PCT     = [3,    16,         38,        43]   # = 100%

TBL1_ROWS = [
    (1,  'n_forbid',                'număr cereri cu decision=forbid',                         'Cereri respinse 403'),
    (2,  'n_events',                'mărimea ferestrei (≤ 20)',                                'Numărul efectiv de evenimente'),
    (3,  'n_secrets',               'număr cereri cu resource=secrets (orice verb)',           'Operații pe secrete'),
    (4,  'n_exec',                  'număr cereri cu sub=exec',                                'Comenzi executate în container'),
    (5,  'n_rbac',                  'număr cereri pe roluri/binding-uri RBAC',                 'Operații RBAC'),
    (6,  'n_create',                'număr verbe create',                                      'Creări de resurse'),
    (7,  'n_delete',                'număr verbe delete',                                      'Ștergeri de resurse'),
    (8,  'n_4xx',                   'număr răspunsuri cu cod HTTP ≥ 400',                      'Erori client'),
    (9,  'n_selfreview',            'număr create pe selfsubject(access|rules)reviews',        'Apeluri kubectl auth can-i (recon permisiuni)'),
    (10, 'n_distinct_resource',     '|{resource}|',                                            'Tipuri distincte de resurse atinse'),
    (11, 'n_distinct_verb',         '|{verb}|',                                                'Verbe distincte'),
    (12, 'n_distinct_ns',           '|{namespace}|',                                           'Namespace-uri distincte'),
    (13, 'secret_ns',               '|{ns : citire de secret}|',                               'Namespace-uri în care s-au citit secrete'),
    (14, 'n_distinct_impersonated', '|{imp : is_imp}|',                                        'Identități distincte impersonate (--as)'),
    (15, 'forbid_ratio',            'n_forbid / n_events',                                     'Procent de respingeri'),
    (16, 'selfreview_ratio',        'n_selfreview / n_events',                                 'Densitatea reconului de permisiuni'),
    (17, 'secret_rate',             'n_secrets / n_events',                                    'Densitatea operațiilor pe secrete'),
    (18, 'rbac_rate',               'n_rbac / n_events',                                       'Densitatea operațiilor RBAC'),
    (19, 'create_rate',             'n_create / n_events',                                     'Densitatea creărilor'),
    (20, 'has_secret',              '1 dacă există citire de secret',                          'Flag prezență citire secret'),
    (21, 'has_exec',                '1 dacă n_exec > 0',                                       'Flag prezență exec'),
    (22, 'has_rbac_write',          '1 dacă există create/update/patch/delete pe RBAC',        'Flag prezență scriere RBAC'),
    (23, 'has_crb',                 '1 dacă există creare ClusterRoleBinding/ClusterRole',     'Flag prezență creare CRB (escaladare)'),
    (24, 'has_forbid',              '1 dacă n_forbid > 0',                                     'Flag prezență respingere'),
    (25, 'has_impersonation',       '1 dacă atacatorul ≠ identitatea efectivă',                'Flag prezență impersonare (T1134)'),
    (26, 'severity',                '3·has_crb + 2·has_exec + 2·(secret_ns≥2) + has_rbac_write + 2·has_impersonation', 'Scor compus al primitivelor rare (0–10)'),
    (27, 'cum_secrets',             'acumulat: citiri secret pe stream-ul actorului',          'Memorie peste ferestre — anti low-and-slow'),
    (28, 'cum_rbac_w',              'acumulat: scrieri RBAC pe stream',                        'Memorie peste ferestre'),
    (29, 'cum_exec',                'acumulat: exec-uri pe stream',                            'Memorie peste ferestre'),
    (30, 'cum_crb',                 'acumulat: creări CRB pe stream',                          'Memorie peste ferestre'),
]

make_adaptive_table(TBL1_HEADERS, TBL1_ROWS, TBL1_PCT)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ============================================================
# TABEL 2 — 8 atribute nefolosite
# ============================================================
h2 = doc.add_heading('Tabelul 2 — Cele 8 atribute nefolosite de clasificator', level=1)
h2.runs[0].font.size = Pt(12)
h2.paragraph_format.space_before = Pt(4)
h2.paragraph_format.space_after = Pt(4)

TBL2_HEADERS = ['#', 'Atribut', 'Rol', 'Motiv']
TBL2_PCT     = [3,    16,         28,    53]   # = 100%

TBL2_ROWS = [
    (31, 'label',             'Etichetă ground-truth (0=benign, 1=atac)',                     'Target predicției, nu input'),
    (32, 'n_list',            'Numărul verbelor list în fereastră',                          'Exclusă din ML — cârjă de densitate (modelul cădea pe atac diluat)'),
    (33, 'n_create_workload', 'Numărul creărilor de Pod/Deployment/Job etc.',                'Exclusă din ML (produce FP pe operatori benigni); folosită de regula hijack'),
    (34, 'has_csr',           'Flag prezență CertificateSigningRequest',                     'Folosită doar de regula persist (allowlist + flag binar = mai robust)'),
    (35, 'has_tokenreq',      'Flag prezență serviceaccounts/token create',                  'Folosită doar de regula persist (TokenRequest abuse)'),
    (36, 'user',              'Identitatea actorului',                                       'Trasabilitate + atribuire etichetă; exclusă din input ca anti-leakage'),
    (37, 'tool',              'Sursa traficului (synthetic/stratus/peirates/rakkess/...)',   'Trasabilitate pentru split tool-disjunct'),
    (38, 'session',           'ID-ul sesiunii de colectare (1–107)',                         'Trasabilitate pentru split sesiune-disjoint și reconstrucția episoadelor'),
]

make_adaptive_table(TBL2_HEADERS, TBL2_ROWS, TBL2_PCT)

doc.save(OUT)
print(f"Saved: {OUT}")
